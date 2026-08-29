"""
Generates the Executive Word Document (.docx) for KBM Functional Scenarios & User Journeys.
"""

import os
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def generate_scenario_diagrams(output_dir):
    os.makedirs(output_dir, exist_ok=True)
    plt.rcParams['font.sans-serif'] = 'Arial'
    plt.rcParams['font.family'] = 'sans-serif'

    # ---------------------------------------------------------
    # Flowchart 1: Tender Request Intake & Validation
    # ---------------------------------------------------------
    fig, ax = plt.subplots(figsize=(11, 6), dpi=300)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 6)
    ax.axis('off')

    ax.text(5.5, 5.6, "Scenario 1: Tender Request Intake & Mandatory Scan Validation Flow", 
            fontsize=13, fontweight='bold', ha='center', va='center', color='#1E3A8A')

    # Channel A
    r_a = patches.FancyBboxPatch((0.5, 3.6), 2.8, 1.3, boxstyle="round,pad=0.06", facecolor="#EFF6FF", edgecolor="#2563EB", lw=1.5)
    ax.add_patch(r_a)
    ax.text(1.9, 4.4, "Channel A: Raslni G2G", fontsize=9.5, fontweight='bold', color="#1E40AF", ha='center')
    ax.text(1.9, 3.9, "Electronic Government\nCorrespondence Message", fontsize=7.5, color="#334155", ha='center')

    # Channel B
    r_b = patches.FancyBboxPatch((0.5, 1.5), 2.8, 1.5, boxstyle="round,pad=0.06", facecolor="#FFFBEB", edgecolor="#D97706", lw=1.5)
    ax.add_patch(r_b)
    ax.text(1.9, 2.6, "Channel B: Paper GM Letter", fontsize=9.5, fontweight='bold', color="#B45309", ha='center')
    ax.text(1.9, 2.0, "Official Physical Letter +\nMandatory Scan (FR-012)", fontsize=7.5, color="#334155", ha='center')

    # Validation Engine
    r_val = patches.FancyBboxPatch((4.2, 2.0), 3.0, 2.4, boxstyle="round,pad=0.06", facecolor="#F0FDF4", edgecolor="#16A34A", lw=1.5)
    ax.add_patch(r_val)
    ax.text(5.7, 4.0, "Intake Validation Engine", fontsize=10, fontweight='bold', color="#15803D", ha='center')
    ax.text(5.7, 3.0, "1. Check Mandatory Scan (PDF)\n2. Validate MoCI Activity Code\n3. Compute SHA-256 Checksum\n4. Assign Sequence Reference", fontsize=7.5, color="#1E293B", ha='center')

    # Destination Stores
    r_db = patches.FancyBboxPatch((8.0, 3.4), 2.5, 1.4, boxstyle="round,pad=0.06", facecolor="#FAF5FF", edgecolor="#7E22CE", lw=1.5)
    ax.add_patch(r_db)
    ax.text(9.25, 4.3, "Azure Cosmos DB", fontsize=9, fontweight='bold', color="#6B21A8", ha='center')
    ax.text(9.25, 3.7, "Tenders Container (/tenantId)\nStatus: INTAKE_PENDING", fontsize=7.5, color="#334155", ha='center')

    r_aud = patches.FancyBboxPatch((8.0, 1.4), 2.5, 1.4, boxstyle="round,pad=0.06", facecolor="#FEF2F2", edgecolor="#DC2626", lw=1.5)
    ax.add_patch(r_aud)
    ax.text(9.25, 2.3, "Cryptographic Ledger", fontsize=9, fontweight='bold', color="#991B1B", ha='center')
    ax.text(9.25, 1.7, "AuditEvents Container\nSHA-256 Hash Chained", fontsize=7.5, color="#334155", ha='center')

    # Arrows
    ax.annotate('', xy=(4.2, 4.0), xytext=(3.3, 4.2), arrowprops=dict(arrowstyle="->", color="#1E3A8A", lw=2))
    ax.annotate('', xy=(4.2, 2.7), xytext=(3.3, 2.3), arrowprops=dict(arrowstyle="->", color="#1E3A8A", lw=2))
    ax.annotate('', xy=(8.0, 4.1), xytext=(7.2, 3.6), arrowprops=dict(arrowstyle="->", color="#16A34A", lw=2))
    ax.annotate('', xy=(8.0, 2.1), xytext=(7.2, 2.6), arrowprops=dict(arrowstyle="->", color="#16A34A", lw=2))

    p1 = os.path.join(output_dir, "scenario1_intake_flow.png")
    plt.tight_layout()
    plt.savefig(p1, dpi=300, bbox_inches='tight')
    plt.close()

    # ---------------------------------------------------------
    # Flowchart 2: Server-Side Eligibility & KNET Purchase
    # ---------------------------------------------------------
    fig, ax = plt.subplots(figsize=(11, 6), dpi=300)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 6)
    ax.axis('off')

    ax.text(5.5, 5.6, "Scenario 3: Vendor Discovery, Server-Side Eligibility & KNET Checkout", 
            fontsize=13, fontweight='bold', ha='center', va='center', color='#1E3A8A')

    boxes_s3 = [
        ("1. Vendor Discovery", "Vendor searches catalog\nGET /api/tenders?tenantId=moi", 1.5, "#EFF6FF", "#1D4ED8"),
        ("2. Server Eligibility", "Activity Overlap Check\n+ 3-Grade Hierarchy Match", 4.1, "#F0FDF4", "#15803D"),
        ("3. Watermark Preview", "Anti-Screenshot Canvas\n(CR + Timestamp + IP)", 6.7, "#FFFBEB", "#B45309"),
        ("4. KNET Checkout", "HMAC SHA-256 Webhook\nReceipt: REC-YYYY-SEQ", 9.3, "#FAF5FF", "#7E22CE")
    ]

    for title, desc, cx, bg_c, brd_c in boxes_s3:
        w, h = 2.2, 3.2
        r = patches.FancyBboxPatch((cx - w/2, 1.5), w, h, boxstyle="round,pad=0.06", facecolor=bg_c, edgecolor=brd_c, lw=1.5)
        ax.add_patch(r)
        ax.text(cx, 4.2, title, fontsize=9.5, fontweight='bold', color=brd_c, ha='center')
        ax.text(cx, 2.8, desc, fontsize=8, color="#1E293B", ha='center', va='center', linespacing=1.4)

    for cx in [2.6, 5.2, 7.8]:
        ax.annotate('', xy=(cx + 0.3, 3.1), xytext=(cx, 3.1), arrowprops=dict(arrowstyle="->", color="#1E3A8A", lw=2.2))

    p2 = os.path.join(output_dir, "scenario2_eligibility_flow.png")
    plt.tight_layout()
    plt.savefig(p2, dpi=300, bbox_inches='tight')
    plt.close()

    print("[SUCCESS] Scenario flowcharts generated.")
    return p1, p2

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9.0)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(30, 41, 59)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def add_callout(doc, title, text, bg_hex="EFF6FF", border_hex="2563EB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    p = cell.paragraphs[0]
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(30, 58, 138)

    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph()

def build_scenarios_document():
    doc_dir = os.path.join(os.path.dirname(__file__), "..", "docs", "scenarios")
    diagrams_dir = os.path.join(doc_dir, "diagrams")
    p1_img, p2_img = generate_scenario_diagrams(diagrams_dir)

    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # ---------------------------------------------------------
    # COVER / HEADER
    # ---------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(28)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("KBM Platform — Functional Scenarios & User Journeys")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run("End-to-End Execution Walkthroughs, Business Validation Rules & Technical Contracts\nAligned with Kuwait Public Tenders Law No. 49/2016 & MOI BRD v0.1")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    add_callout(doc, "Scenario Walkthrough Metadata",
                "• Document Reference: KBM-FSD-SCENARIOS-2026\n"
                "• Scope: 4 Core End-to-End Business Scenarios\n"
                "• Target Systems: Staff Portal, Vendor Portal, Tenant Admin Portal, KNET Gateway, Azure Cosmos DB\n"
                "• Date of Release: August 2026 | Status: Approved")

    # ---------------------------------------------------------
    # SCENARIO 1: TENDER INTAKE
    # ---------------------------------------------------------
    h1 = doc.add_heading("1. Scenario 1: Tender Request Intake & Registration Cycle", level=1)
    for r in h1.runs: r.font.name = "Arial"; r.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "This scenario covers the end-to-end process of originating, validating, and registering procurement requests from government sectors. "
        "The platform supports both automated electronic intake via Raslni G2G (FR-002) and paper General Manager letters with mandatory scanned attachments (FR-012)."
    )

    if os.path.exists(p1_img):
        doc.add_picture(p1_img, width=Inches(6.5))
        p_cap = doc.add_paragraph("Figure 1: Scenario 1 - Tender Request Intake & Mandatory Scan Validation Flow")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p_cap.runs: r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    s1_headers = ["Step", "Actor", "Action / Screen", "Business Validation Rule", "Cosmos DB State Change"]
    s1_data = [
        ["1.1", "Requesting Officer", "Staff Portal: New Intake Screen", "Selects Channel: Paper Letter or Raslni G2G.", "Form initialized with tenant context."],
        ["1.2", "Requesting Officer", "Fills Title, Budget, Attaches Scanned GM Letter", "FR-012 Mandatory Check: If scanned letter missing, submit is rejected with HTTP 400.", "Blob: kbmbootstrapsa/scans/req-xxxx.pdf"],
        ["1.3", "Procurement Engine", "POST /api/requests", "MIME validation (PDF/JPEG allowlist), SHA-256 payload checksum calculation.", "Tenders container: status: INTAKE_PENDING"],
        ["1.4", "Audit Store", "Append Audit Event", "SHA-256 Hash Chaining (previousHash -> hash).", "AuditEvents container: new block appended"],
        ["1.5", "Tender Committee Staff", "Intake Review & Classification", "Selects MoCI activity codes (IT-SYS-01) and sets grade rule (SECOND_AND_ABOVE).", "Tenders container: status: PUBLISHED"]
    ]
    create_styled_table(doc, s1_headers, s1_data, [0.7, 1.3, 1.6, 1.8, 1.4])

    # ---------------------------------------------------------
    # SCENARIO 2: VENDOR REGISTRATION & GRADING
    # ---------------------------------------------------------
    h2 = doc.add_heading("2. Scenario 2: Vendor Registration, 3-Grade Classification & Subscription", level=1)
    for r in h2.runs: r.font.name = "Arial"; r.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "Commercial enterprises apply through the Vendor Portal by submitting commercial licenses and MoCI activity registrations. "
        "The Tenant Admin evaluates the technical capacity and assigns a grade (First, Second, or Third Grade). The vendor then activates an annual 1-year KNET subscription."
    )

    s2_headers = ["Step", "Phase", "Operational Description", "Statutory Rule / BRD Reference", "Cosmos DB Outcome"]
    s2_data = [
        ["2.1", "Registration", "Vendor submits company profile, CR number, civil IDs, and commercial activity codes.", "FR-015: Online registration with dossier upload.", "Vendors: status: PENDING_APPROVAL"],
        ["2.2", "Grading Evaluation", "Tenant Admin evaluates past projects and capital, assigning Grade FIRST, SECOND, or THIRD.", "FR-009 & FR-016: 3-tier classification hierarchy.", "Vendors: grade: FIRST, status: APPROVED"],
        ["2.3", "KNET Subscription", "Vendor completes 1-year subscription checkout via KNET (100 KWD/yr).", "FR-017: 1-Year KNET subscription validity.", "Vendors: subscriptionExpiresAt: +365d"],
        ["2.4", "Fee Exemption", "Admin can grant fee exemptions to designated SME entities.", "FR-018: Fee exemption bypasses checkout.", "Vendors: isFeeExempt: true"],
        ["2.5", "Suspension & Auto-Restore", "Admin sets temporary date-range block (blockedUntil). System automatically restores access on expiry.", "FR-013 & FR-014: Date-range suspension with auto-reinstatement.", "Vendors: isBlocked: false when now > blockedUntil"]
    ]
    create_styled_table(doc, s2_headers, s2_data, [0.7, 1.3, 1.8, 1.6, 1.4])

    # ---------------------------------------------------------
    # SCENARIO 3: TENDER ELIGIBILITY & KNET PURCHASE
    # ---------------------------------------------------------
    h3 = doc.add_heading("3. Scenario 3: Tender Discovery, Server-Side Eligibility & KNET Purchase", level=1)
    for r in h3.runs: r.font.name = "Arial"; r.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "Vendors discover tenders matching their registered commercial activities. The platform enforces server-side eligibility checks (FR-020 & SEC-006), "
        "displays dynamic anti-screenshot canvas watermarking (SEC-001), and generates official branded purchase receipts (DOC-002) upon successful KNET checkout."
    )

    if os.path.exists(p2_img):
        doc.add_picture(p2_img, width=Inches(6.5))
        p_cap = doc.add_paragraph("Figure 2: Scenario 3 - Vendor Discovery, Server-Side Eligibility & KNET Checkout Flow")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p_cap.runs: r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    s3_headers = ["Vendor Grade", "Tender Target Grade", "Matching Rule", "Eligibility Result", "System Enforcement"]
    s3_data = [
        ["First Grade (FIRST)", "Second Grade (SECOND)", "GRADE_AND_ABOVE", "ELIGIBLE (PASS)", "Full booklet preview enabled; KNET purchase active."],
        ["Third Grade (THIRD)", "Second Grade (SECOND)", "GRADE_AND_ABOVE", "INELIGIBLE (BLOCKED)", "HTTP 403: Vendor grade does not meet minimum threshold."],
        ["First Grade (FIRST)", "Second Grade (SECOND)", "EXACT_GRADE", "INELIGIBLE (BLOCKED)", "HTTP 403: Exact grade match required for targeted tender."],
        ["Any Grade", "Any Grade", "No Activity Overlap", "INELIGIBLE (HIDDEN)", "Tender hidden from catalog; direct API access rejected."]
    ]
    create_styled_table(doc, s3_headers, s3_data, [1.4, 1.4, 1.4, 1.4, 1.4])

    # ---------------------------------------------------------
    # SCENARIO 4: STATUTORY WORKFLOW & STATE AUDIT OVERSIGHT
    # ---------------------------------------------------------
    h4 = doc.add_heading("4. Scenario 4: Statutory Approval Workflow & State Audit Bureau Oversight", level=1)
    for r in h4.runs: r.font.name = "Arial"; r.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "KBM manages the 35-step statutory approval state machine for public tenders, tracking SLA deadlines, return-for-correction loops, "
        "and pre-contract statutory oversight by the State Audit Bureau (ديوان المحاسبة) before Undersecretary contract signing."
    )

    s4_headers = ["Workflow Milestone", "Responsible Authority", "Statutory Task Action", "SLA Threshold", "Regulatory Output"]
    s4_data = [
        ["Milestone 1: Intake & CAIT", "Requesting Dept & CAIT", "Needs survey, cost estimation, project code assignment.", "5 Business Days", "CAIT Portal Entry Number"],
        ["Milestone 2: Legal Review", "Fatwa & Legislation Dept", "Legal review of tender conditions and penalty clauses.", "10 Business Days", "Official Legal Review Seal"],
        ["Milestone 3: CAPT Approval", "Central Agency for Public Tenders", "Board authorization and Kuwait Alyawm gazette publication.", "14 Business Days", "Official Gazette Notice Ref"],
        ["Milestone 4: Pre-Audit Oversight", "State Audit Bureau (ديوان المحاسبة)", "Statutory pre-contract audit review (الرقابة المسبقة).", "15 Business Days", "Pre-Audit Approval Letter"],
        ["Milestone 5: Contract Execution", "Undersecretary (وكيل الوزارة)", "Final contract signing 100% and bank guarantee deposit.", "3 Business Days", "Cosmos DB SHA-256 Audit Seal"]
    ]
    create_styled_table(doc, s4_headers, s4_data, [1.4, 1.5, 1.7, 1.1, 1.3])

    # Save documents
    final_path = os.path.join(doc_dir, "KBM_Functional_Scenarios_and_User_Journeys.docx")
    doc.save(final_path)

    root_path = os.path.join(os.path.dirname(__file__), "..", "KBM_Functional_Scenarios_and_User_Journeys.docx")
    doc.save(root_path)

    print(f"[SUCCESS] Functional Scenarios Word document generated at: {final_path}")
    print(f"[SUCCESS] Root copy generated at: {root_path}")

if __name__ == "__main__":
    build_scenarios_document()
