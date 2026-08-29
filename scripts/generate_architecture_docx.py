"""
Generates a comprehensive executive Word document (.docx) for the KBM Platform Architecture.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=160, right=160)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(30, 41, 59)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph() # Spacing
    return table

def add_callout(doc, title, text, bg_hex="EFF6FF", border_hex="2563EB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    p = cell.paragraphs[0]
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(30, 58, 138)

    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph()

def build_document():
    doc = docx.Document()

    # Page Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # -------------------------------------------------------------
    # 1. COVER / TITLE
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("KBM Procurement & Tender Management Platform")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Complete System Architecture & Technical Specification Document\nEnterprise Multi-Tenant SaaS, Zero-Cost Bootstrap & Cloud Infrastructure")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    add_callout(doc, "Document Governance & Release Metadata",
                "• Version: 1.0 (Production Architecture Baseline)\n"
                "• Product: KBM Multi-Tenant Procurement SaaS Platform\n"
                "• Regulatory Target: State of Kuwait Public Tenders Law No. 49/2016\n"
                "• Target Cloud: Microsoft Azure (Enterprise Multi-Tenant / Commercial Marketplace)\n"
                "• Date of Issue: August 2026 | Classification: Confidential / Executive")

    # -------------------------------------------------------------
    # 2. EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    h1 = doc.add_heading("1. Executive Summary & Architectural Vision", level=1)
    h1.paragraph_format.space_before = Pt(18)
    for run in h1.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 58, 138)

    p = doc.add_paragraph(
        "The KBM Platform is an enterprise-grade, multi-tenant Software-as-a-Service (SaaS) solution engineered "
        "to modernize and digitize public procurement, practices (الممارسات), and tenders (المناقصات) in strict "
        "compliance with State of Kuwait Public Tenders Law (Law No. 49/2016) and relevant regulatory bodies, including "
        "the Central Agency for Public Tenders (CAPT), Fatwa & Legislation Department, the State Audit Bureau (ديوان المحاسبة), "
        "and the Central Agency for Information Technology (CAIT)."
    )
    p.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph(
        "The architecture is constructed around a Zero-Cost Bootstrap Profile ($0.00/month on Azure Free Tier) "
        "for instant client validation and demo scenarios, seamlessly upgradeable to a High-Availability Enterprise Production Profile "
        "featuring Azure Cosmos DB NoSQL partitioning, Azure Key Vault Hardware Security Module (HSM) encryption, and Microsoft Commercial Marketplace SaaS Fulfillment v2 integration."
    )
    p2.paragraph_format.space_after = Pt(14)

    # -------------------------------------------------------------
    # 3. HIGH-LEVEL TOPOLOGY & TIERS
    # -------------------------------------------------------------
    h2 = doc.add_heading("2. High-Level Architectural Topology", level=1)
    h2.paragraph_format.space_before = Pt(18)
    for run in h2.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "The KBM system topology is structured into 6 decoupled architectural layers designed for maximum modularity, "
        "multi-tenant data isolation, zero bundle bloat, and enterprise compliance:"
    )

    topology_headers = ["Layer / Tier", "Primary Components", "Architectural Responsibilities"]
    topology_data = [
        ["1. Client Presentation", "Vanilla HTML5 / CSS3 / ES2022, Canvas API, i18n Engine", "Role-adapted portals (Staff, Vendor, Admin, Marketplace SaaS), 5-stage interactive lifecycle journey visualizer, dynamic anti-screenshot canvas watermarking."],
        ["2. Edge & Gateway", "Node.js Edge Gateway, ASVS Security Headers, Rate Limiter", "Reverse proxy on port 3000, OWASP ASVS headers (HSTS, CSP, X-Frame-Options DENY), sliding-window rate limiting (200 req/min), tenant policy boundary enforcement."],
        ["3. Core Microservices", "Procurement, Workflow, Vendor, Payment, Document, Audit, Marketplace", "Domain-driven bounded contexts managing 35-step statutory workflows, MoCI activity & 3-grade vendor eligibility, KNET checkout, and SaaS fulfillment."],
        ["4. Integration Adapters", "Raslni G2G, MoCI ISIC Catalog, Entra ID SSO Adapter", "Inter-ministerial correspondence ingestion, Ministry of Commerce activity codes taxonomy, and Active Directory security group-to-role mappings."],
        ["5. Data & Crypto Ledger", "Azure Cosmos DB, Azure Blob / Azurite, Azure Key Vault", "Partitioned document persistence (/tenantId), append-only SHA-256 hash chaining audit ledger, encrypted specification storage, and HSM key management."],
        ["6. External Systems", "Kuwait G2G, MoCI, KNET Payment Gateway, Azure Marketplace", "Government correspondence network, electronic banking checkout, and Microsoft AppSource transactable subscription provisioning."]
    ]
    create_styled_table(doc, topology_headers, topology_data, [1.5, 2.0, 3.0])

    # -------------------------------------------------------------
    # 4. TENDER LIFECYCLE & STATUTORY APPROVAL PIPELINE
    # -------------------------------------------------------------
    h3 = doc.add_heading("3. End-to-End Tender Lifecycle & Regulatory Journey", level=1)
    h3.paragraph_format.space_before = Pt(18)
    for run in h3.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "KBM digitalizes the complete statutory workflow of Kuwait public procurement, tracking SLA deadlines, "
        "responsible government bodies, and mandatory attachments across 5 primary stages:"
    )

    stages_headers = ["Stage", "Stage Name (Bilingual)", "Responsible Oversight Body", "Mandatory Statutory Actions & Governance"]
    stages_data = [
        ["Stage 1", "1. Preparation & Budget\nالتجهيز والميزانية", "Systems Sector & Ministry of Finance\nالنظم & المالية", "Needs assessment survey, cost estimation, CAIT portal project entry, Ministry of Finance budget allocation approval."],
        ["Stage 2", "2. Technical & Legal Review\nالفتوى والتشريع ولجنة الشراء", "Purchase Committee & Legal Department\nلجنة الشراء & إدارة الفتوى", "Purchase committee technical review, terms booklet finalization, Fatwa & Legislation legal review, internal audit seal."],
        ["Stage 3", "3. Public Launch & CAPT\nإعلان الطرح والتأهيل", "Central Agency for Public Tenders (CAPT)\nالجهاز المركزي للمناقصات", "CAPT formal approval, official gazette (Kuwait Alyawm) publishing, server-side MoCI activity and grading eligibility filter."],
        ["Stage 4", "4. Bid Opening & Award\nدراسة العطاءات والترسية", "Bid Opening & Technical Evaluation Comm.\nلجنة فتح المظاريف والدراسة الفنية", "Dynamic anti-screenshot booklet access, KNET purchase checkout, financial/technical bid scoring, final award decision."],
        ["Stage 5", "5. State Audit & Contract (100%)\nديوان المحاسبة وتوقيع العقد", "State Audit Bureau & Undersecretary\nديوان المحاسبة & وكيل الوزارة", "State Audit Bureau pre-contract approval (الرقابة المسبقة), final performance bond submission, Undersecretary contract execution 100%."]
    ]
    create_styled_table(doc, stages_headers, stages_data, [1.0, 1.8, 1.7, 2.0])

    # -------------------------------------------------------------
    # 5. CORE DOMAIN MICROSERVICES
    # -------------------------------------------------------------
    h4 = doc.add_heading("4. Microservices & Bounded Contexts Specification", level=1)
    h4.paragraph_format.space_before = Pt(18)
    for run in h4.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "The backend is partitioned into discrete, independent domain services adhering to Domain-Driven Design (DDD) principles:"
    )

    services_headers = ["Microservice Repository", "Core Domain Responsibility", "Key APIs & Data Structures", "Cryptographic / Governance Control"]
    services_data = [
        ["kbm-platform-workflow-service", "State Machine & SLA Orchestration", "POST /api/workflow/tasks/:id/transition\nGET /api/workflow/instances/:id", "Versioned templates (8-step Practices, 35-step Tenders), return-for-correction loops, role escalation."],
        ["kbm-platform-procurement-service", "Intake & Server-Side Eligibility", "POST /api/requests\nPOST /api/tenders\nGET /api/tenders", "Validates mandatory scanned GM letters (FR-012), strict server-side MoCI activity & grade filtering."],
        ["kbm-platform-vendor-service", "Vendor Registry & Governance", "POST /api/vendors/register\nPOST /api/vendors/:id/grade\nPOST /api/vendors/:id/block", "3-tier classification (First, Second, Third), date-range suspensions with auto-reinstatement, fee exemptions."],
        ["kbm-platform-payment-service", "KNET Checkout & Receipts", "POST /api/payments/checkout\nPOST /api/payments/knet/callback", "Idempotent payment callback processing, HMAC SHA-256 signatures, tenant-branded receipts (REC-YYYY-SEQ)."],
        ["kbm-platform-document-service", "Document Ingestion & Watermarking", "POST /api/documents\nGET /api/documents/:id/watermark", "MIME type allowlisting, SHA-256 upload checksums, dynamic anti-screenshot canvas watermark schema."],
        ["kbm-platform-audit-service", "Immutable Compliance Ledger", "POST /api/audit/events\nGET /api/audit/verify\nGET /api/audit/export", "Append-only SHA-256 hash chaining (previousHash -> hash), automated mathematical tamper detection."],
        ["kbm-platform-marketplace-service", "Microsoft Marketplace SaaS", "POST /api/marketplace/resolve\nPOST /api/marketplace/activate\nPOST /api/marketplace/webhook", "Turnkey SaaS Fulfillment API v2 integration, token resolution, subscription plan lifecycle webhooks."]
    ]
    create_styled_table(doc, services_headers, services_data, [1.6, 1.8, 1.8, 1.3])

    # -------------------------------------------------------------
    # 6. SECURITY & CRYPTOGRAPHIC ARCHITECTURE
    # -------------------------------------------------------------
    h5 = doc.add_heading("5. Security Architecture, Cryptography & Threat Mitigation", level=1)
    h5.paragraph_format.space_before = Pt(18)
    for run in h5.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "KBM implements a defense-in-depth security model across authentication, authorization, cryptographic data protection, and edge routing:"
    )

    sec_headers = ["Security Vector", "OWASP / STRIDE Threat", "Architecture Defense", "Code Implementation & Tests"]
    sec_data = [
        ["Tenant Isolation", "Elevation of Privilege / BOLA", "Multi-tenant boundary evaluation on every direct object access; cross-tenant requests strictly forbidden.", "services/shared/tenant-policy.js\ntenant-policy.test.js"],
        ["Audit Trail Immutability", "Tampering / Repudiation", "Append-only event store with SHA-256 cryptographic hash chaining linking each event to the predecessor.", "kbm-platform-audit-service/audit-store.js\naudit-store.test.js"],
        ["Document Protection", "Information Disclosure / Leakage", "MIME allowlisting, SHA-256 checksums, and dynamic canvas watermarking embedding recipient CR & IP.", "kbm-platform-document-service/document-manager.js\ndocument-manager.test.js"],
        ["Payment Webhooks", "Tampering / Forgery", "Idempotent transaction tracking and HMAC-SHA256 signature verification over payment ID and amounts.", "kbm-platform-payment-service/payment-manager.js\npayment-manager.test.js"],
        ["Edge Security", "Denial of Service / XSS", "OWASP ASVS headers (HSTS, CSP, X-Frame-Options: DENY, nosniff) and sliding-window rate limiting (200 req/min).", "kbm-platform-edge/edge-gateway.js\nedge-gateway.test.js"],
        ["Path Traversal", "Insecure File Access", "Strict root directory prefix verification preventing directory traversal (../) attacks on static assets.", "services/web/web-server.js\nproxy-router.test.js"]
    ]
    create_styled_table(doc, sec_headers, sec_data, [1.4, 1.6, 2.0, 1.5])

    # -------------------------------------------------------------
    # 7. DATABASE & PERSISTENCE ARCHITECTURE
    # -------------------------------------------------------------
    h6 = doc.add_heading("6. Database & Persistence Architecture (Azure Cosmos DB Free Tier)", level=1)
    h6.paragraph_format.space_before = Pt(18)
    for run in h6.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "The KBM platform utilizes Azure Cosmos DB Free Tier (enableFreeTier: true) providing 1,000 RU/s throughput "
        "and 25 GB storage free for the lifetime of the Azure subscription at $0.00/month:"
    )

    db_headers = ["Cosmos DB Container", "Partition Key", "Data Managed & Stored", "Cryptographic / Consistency Model"]
    db_data = [
        ["Tenders", "/tenantId", "Active tenders, published practices, MoCI commercial activity codes, prices, and eligibility rules.", "Session Consistency, Indexed on /activities and /gradeRule"],
        ["Vendors", "/tenantId", "Vendor company dossiers, 3-grade classifications, temporary suspensions, and fee exemptions.", "Session Consistency, Unique Commercial Registration indexing"],
        ["Workflows", "/tenantId", "35-step statutory approval task instances, SLA trackers, return-for-correction histories.", "Session Consistency, State machine transition audit links"],
        ["AuditEvents", "/tenantId", "Append-only immutable compliance log linking events with SHA-256 hash chaining (previousHash -> hash).", "Cryptographic Tamper-Evident Ledger, Mathematical verification"]
    ]
    create_styled_table(doc, db_headers, db_data, [1.5, 1.3, 2.3, 1.7])

    add_callout(doc, "Live Azure Cosmos DB Free Tier Provisioning Details",
                "• Account: kbm-cosmos-uyhsofjy5a23s.documents.azure.com (West Europe)\n"
                "• Database: kbm-procurement-db (Shared 1,000 RU/s Free Tier)\n"
                "• Partition Key: /tenantId (Physical & Logical isolation across all 4 containers)\n"
                "• Cost: $0.00 / month (100% Lifetime Azure Free Tier)")

    # -------------------------------------------------------------
    # 8. CLOUD INFRASTRUCTURE & CI/CD DEPLOYMENT
    # -------------------------------------------------------------
    h7 = doc.add_heading("7. Cloud Infrastructure as Code (IaC) & CI/CD Pipelines", level=1)
    h7.paragraph_format.space_before = Pt(18)
    for run in h7.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "Infrastructure is fully declared via Azure Bicep (kbm-platform-infrastructure/main.bicep), supporting automated "
        "deployments through PowerShell scripts and GitHub Actions workflows:"
    )

    infra_headers = ["Deployment Component", "Configuration & Implementation", "Purpose & Automation Role"]
    infra_data = [
        ["Azure Bicep (main.bicep)", "Parameterized templates (deploymentProfile: bootstrap | production)", "Automates resource provisioning for App Service, Storage, Key Vault, and Cosmos DB."],
        ["PowerShell Script (deploy-azure.ps1)", "Native Azure CLI orchestrator with auto-PATH discovery", "One-click deployment script provisioning resource groups, Bicep templates, and ZIP deployments."],
        ["GitHub Actions CI (.github/workflows/ci.yml)", "Automated testing on Node 20.x & 22.x across push and pull requests", "Executes bootstrap governance validation, 35 unit/integration tests, and E2E cycle verification."],
        ["GitHub Actions Deploy (.github/workflows/deploy-azure.yml)", "Automated CD deployment to Azure App Service / Container Apps", "Deploys code upon push to main/master branches or on-demand workflow dispatch."]
    ]
    create_styled_table(doc, infra_headers, infra_data, [2.0, 2.3, 2.2])

    # Save Word Document
    output_dir = os.path.join(os.path.dirname(__file__), "..", "docs", "architecture")
    os.makedirs(output_dir, exist_ok=True)
    docx_path = os.path.join(output_dir, "KBM_Complete_System_Architecture.docx")
    doc.save(docx_path)

    # Also save to project root for convenience
    root_docx_path = os.path.join(os.path.dirname(__file__), "..", "KBM_Complete_System_Architecture.docx")
    doc.save(root_docx_path)

    print(f"[SUCCESS] Word document generated at: {docx_path}")
    print(f"[SUCCESS] Copy generated at root: {root_docx_path}")

if __name__ == "__main__":
    build_document()
